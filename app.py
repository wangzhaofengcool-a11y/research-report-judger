import streamlit as st, time, json, re, unicodedata, io
from datetime import datetime
from urllib.request import Request, urlopen
from urllib.error import HTTPError
from pathlib import Path

st.set_page_config(page_title="研究报告中判官", page_icon="📋", layout="wide")
st.markdown("""<style>
@media(max-width:768px){button{min-height:44px!important;width:100%!important}}
.report-box{max-height:70vh;overflow-y:auto;padding:1.5rem;border:1px solid #e0e0e0;border-radius:10px;background:#fafbfc;font-size:.92rem;line-height:1.75}
.stButton>button{border-radius:8px}
.history-item{border:1px solid #e0e0e0;border-radius:8px;padding:10px;margin:5px 0;cursor:pointer}
</style>""", unsafe_allow_html=True)

for k,v in {"running":False,"done":False,"reports":[],"all_results":[],"files":[],"search":{},"history":[]}.items():
    if k not in st.session_state: st.session_state[k]=v
def reset():
    for k in ["running","done","reports","all_results","files","search"]:
        st.session_state[k]=False if k in("running","done")else([]if k in("reports","all_results","files")else({}if k=="search"else""))
def sk(key,default=""):
    try:return st.secrets.get(key,default)
    except:
        import os;return os.environ.get(key,default)

# Parsers
def parse_pdf(fb,fn):
    import fitz;doc=None
    try:
        doc=fitz.open(stream=fb,filetype="pdf")
        if doc.is_encrypted:
            try:doc.authenticate("")
            except:return{"fn":fn,"txt":"","chars":0,"err":"PDF加密"}
        total="";empty=0;imgs=0
        for page in doc:
            text=page.get_text("text")
            if not text.strip():
                blocks=page.get_text("blocks")
                text="\n".join(b[4] for b in blocks if len(b)>4 and b[6]==0)
            if not text.strip():
                d=page.get_text("dict")
                for block in d.get("blocks",[]):
                    if block.get("type")==0:
                        for line in block.get("lines",[]):
                            for span in line.get("spans",[]):
                                if span.get("text","").strip():text+=span["text"]+" "
                    elif block.get("type")==1:imgs+=1
            if not text.strip():empty+=1
            total+=text+"\n"
        doc.close()
        if len(total)<100 and empty>doc.page_count*0.5:
            return{"fn":fn,"txt":"","chars":0,"err":f"含{imgs}张图片文本不足"}
        return{"fn":fn,"txt":total,"chars":len(total),"pages":doc.page_count}
    except Exception as e:
        try:doc.close()
        except:pass
        return{"fn":fn,"txt":"","chars":0,"err":str(e)[:80]}

def parse_docx(fb,fn):
    from docx import Document
    try:
        doc=Document(io.BytesIO(fb))
        paras=[p.text for p in doc.paragraphs if p.text.strip()]
        for t in doc.tables:
            for r in t.rows:
                rt=" | ".join(c.text for c in r.cells if c.text.strip())
                if rt.strip():paras.append(rt)
        txt="\n\n".join(paras)
        return{"fn":fn,"txt":txt,"chars":len(txt)}
    except Exception as e:return{"fn":fn,"txt":"","chars":0,"err":str(e)[:80]}

def parse_html(fb,fn):
    from bs4 import BeautifulSoup
    try:
        soup=BeautifulSoup(fb,"html.parser")
        for tag in soup(["script","style","nav","footer","img","svg"]):tag.decompose()
        return{"fn":fn,"txt":soup.get_text(separator="\n",strip=True),"chars":len(soup.get_text(separator="\n",strip=True))}
    except Exception as e:return{"fn":fn,"txt":"","chars":0,"err":str(e)[:80]}

def parse_text(fb,fn):
    for enc in["utf-8","gbk","gb2312","gb18030","latin-1"]:
        try:return{"fn":fn,"txt":fb.decode(enc),"chars":len(fb.decode(enc))}
        except:continue
    return{"fn":fn,"txt":"","chars":0,"err":"编码错误"}

def parse_doc(fb,fn):
    ext=Path(fn).suffix.lower()
    if ext==".pdf":return parse_pdf(fb,fn)
    if ext in(".docx",".doc"):return parse_docx(fb,fn)
    if ext in(".html",".htm"):return parse_html(fb,fn)
    return parse_text(fb,fn)

# LLM
def call_llm(pid,pk,pm,sys_msg,usr_msg,max_tok=4096):
    if pid in("deepseek","glm"):
        base="https://api.deepseek.com" if pid=="deepseek" else "https://open.bigmodel.cn/api/paas/v4"
        url=f"{base}/chat/completions"
        body={"model":pm,"messages":[{"role":"system","content":sys_msg},{"role":"user","content":usr_msg}],"temperature":0.3,"max_tokens":max_tok}
        h={"Authorization":f"Bearer {pk}","Content-Type":"application/json"}
    else:
        url=f"https://generativelanguage.googleapis.com/v1beta/models/{pm}:generateContent?key={pk}"
        body={"system_instruction":{"parts":[{"text":sys_msg}]},"contents":[{"role":"user","parts":[{"text":usr_msg}]}],"generationConfig":{"temperature":0.3,"maxOutputTokens":max_tok}}
        r=urlopen(Request(url,data=json.dumps(body).encode(),headers={"Content-Type":"application/json"},method="POST"),timeout=300)
        return json.loads(r.read().decode())["candidates"][0]["content"]["parts"][0]["text"]
    r=urlopen(Request(url,data=json.dumps(body).encode(),headers=h,method="POST"),timeout=300)
    return json.loads(r.read().decode())["choices"][0]["message"]["content"]

# Web search with dual strategy
def search_web(query,n=3):
    results=[]
    # Strategy 1: DuckDuckGo Instant Answer API
    try:
        r=urlopen(Request(f"https://api.duckduckgo.com/?q={query}&format=json&no_html=1&skip_disambig=1"),timeout=10)
        d=json.loads(r.read().decode())
        for t in d.get("RelatedTopics",[])[:n]:
            if isinstance(t,dict)and"Text"in t:results.append({"t":t.get("FirstURL",""),"s":t.get("Text","")[:200]})
        if d.get("AbstractText"):results.append({"t":d.get("AbstractSource",""),"s":d.get("AbstractText","")[:200]})
    except:pass

    # Strategy 2: DDG HTML search as fallback
    if len(results)<2:
        try:
            r=urlopen(Request(f"https://html.duckduckgo.com/html/?q={query}"),timeout=10)
            html=r.read().decode("utf-8",errors="ignore")
            links=re.findall(r'<a[^>]*class="result__a"[^>]*href="([^"]*)"[^>]*>(.*?)</a>',html)
            snippets=re.findall(r'<a[^>]*class="result__snippet"[^>]*>(.*?)</a>',html)
            for i,(href,title) in enumerate(links[:n]):
                s=snippets[i] if i<len(snippets) else ""
                s=re.sub(r'<[^>]+>','',s)[:200]
                results.append({"t":title.strip(),"s":s})
        except:pass
    return results[:n]

def extract_json(txt):
    try:return json.loads(txt)
    except:pass
    m=re.search(r'```(?:json)?\s*([\s\S]*?)```',txt)
    if m:
        try:return json.loads(m.group(1))
        except:pass
    m=re.search(r'\{[\s\S]*\}',txt)
    if m:
        try:return json.loads(m.group(0))
        except:pass
    return{}

# ── Startup ──
st.title("研究报告中判官")
ds_key=sk("DEEPSEEK_API_KEY");ds_model=sk("DEEPSEEK_MODEL","deepseek-chat")
gk=sk("GOOGLE_API_KEY");glk=sk("GLM_API_KEY")
has_keys=bool(ds_key or gk or glk)

if not has_keys:
    st.warning("请在 Settings > Secrets 中配置 API Key")
    st.code("DEEPSEEK_API_KEY=\"sk-xxx\"\nGOOGLE_API_KEY=\"xxx\"\nGLM_API_KEY=\"xxx\"")
else:
    # Sidebar
    with st.sidebar:
        st.header("模型");models=[]
        if ds_key:st.checkbox("DeepSeek V4 Pro",value=True,disabled=True,key="mds");models.append(("deepseek",ds_key,ds_model))
        if gk:
            if st.checkbox("Gemini 3.5 Flash",value=False,key="mgg"):models.append(("google",gk,sk("GOOGLE_MODEL","gemini-2.5-flash")))
        if glk:
            if st.checkbox("GLM 5.1 Pro",value=False,key="mglm"):models.append(("glm",glk,sk("GLM_MODEL","glm-4.5")))
        st.caption(f"已选{len(models)}个")
        st.markdown("---")
        st.subheader("历史记录")
        if st.button("清空历史",use_container_width=True):st.session_state.history=[]
        st.caption(f"共{len(st.session_state.history)}条")

    # Main
    st.caption("上传文档→多模型交叉验证→修正结论报告");st.markdown("---")
    st.subheader("上传文档")
    up=st.file_uploader("PDF/DOCX/MD/TXT/HTML|≤50MB|≤200MB总",type=["pdf","docx","md","txt","html","htm"],accept_multiple_files=True)
    if up:
        ok=True;tmb=sum(f.size for f in up)/1024/1024
        for f in up:
            if f.size>50*1024*1024:st.error(f"{f.name}超50MB");ok=False
        if tmb>200:st.error(f"总{tmb:.1f}MB超限");ok=False
        if ok:st.success(f"{len(up)}个文件({tmb:.1f}MB)")
    mode=st.radio("处理模式",["逐一分析（每篇独立报告）","合并分析（综合结论）"],horizontal=True,key="mode")
    pm="separate" if mode=="逐一分析" else "merged"
    ci=st.text_area("自定义审校要求（可选）",placeholder="关注第三章方法论\n对比2024年后数据",height=80,key="ci")
    c1,c2=st.columns([3,1])
    with c1:go=st.button("开始审校分析",type="primary",disabled=not up or st.session_state.running,use_container_width=True)
    with c2:
        if st.button("重置",use_container_width=True):reset();st.rerun()
    if go:
        st.session_state.running=True;st.session_state.done=False
        st.session_state.files=[(f.name,f.read())for f in up]
        st.session_state.reports=[];st.session_state.all_results=[];st.session_state.search={}
        st.rerun()

# ── History display ──
if st.session_state.history and not st.session_state.running:
    with st.expander(f"📚 历史记录（{len(st.session_state.history)}条）"):
        for i,h in enumerate(reversed(st.session_state.history)):
            if st.button(f"{h['time']} — {h['title']}",key=f"hist_{i}"):
                st.session_state.reports=[h['report']];st.session_state.all_results=h.get('results',[]);st.session_state.done=True
                st.rerun()

# ── Analysis Pipeline ──
if st.session_state.running and not st.session_state.done and has_keys:
    st.markdown("---");st.subheader("分析进度")
    pbar=st.progress(0);stat=st.empty();det=st.expander("解析详情",expanded=True)
    try:
        stat.info("解析文档...");pbar.progress(3);parsed=[]
        with det:
            for i,(fn,fb) in enumerate(st.session_state.files):
                d=parse_doc(fb,fn);pbar.progress(3+int(9*(i+1)/len(st.session_state.files)))
                if"err"in d:st.warning(f"{fn}:{d['err']}")
                else:parsed.append(d);st.caption(f"{fn}→{d['chars']:,}字符"+(f",{d.get('pages',0)}页"if d.get('pages')else""))
        if not parsed:st.error("解析失败");st.session_state.running=False;st.stop()
        pbar.progress(12)

        # Sanitize all docs
        for d in parsed:
            d["txt"]="".join(c for c in d["txt"] if c.isprintable()or c in"\n\r\t")
            d["txt"]=unicodedata.normalize("NFKC",d["txt"])

        all_reports=[];all_results=[]

        # Determine documents to process
        if pm=="merged":
            combined=""
            for i,d in enumerate(parsed):combined+=f"\n\n=====文档{i+1}:{d['fn']}=====\n\n{d['txt']}"
            doc_groups=[({"fn":"合并文档","txt":combined,"chars":len(combined)},combined)]
        else:
            doc_groups=[(d,d["txt"]) for d in parsed]

        total_groups=len(doc_groups)
        for gidx,(doc_info,doc_text) in enumerate(doc_groups):
            st.caption(f"--- 分析 {gidx+1}/{total_groups}: {doc_info['fn']} ({len(doc_text):,}字符) ---")
            base_progress=12+int(83*(gidx/total_groups))
            group_progress=83//total_groups

            # Step 1: Extract claims
            stat.info(f"DeepSeek 提取结论...({gidx+1}/{total_groups})");pbar.progress(base_progress+int(group_progress*0.2))
            S1="你是学术审校专家。提取所有核心结论、分论点、论据。仅输出JSON:{\"conclusions\":[{\"text\":\"...\"}],\"arguments\":[{\"text\":\"...\",\"parent\":\"...\"}],\"evidence\":[{\"text\":\"...\",\"type\":\"数据/引用/推理\",\"supports\":\"...\"}]}"
            U1=f"文档：\n\n{doc_text}"+(f"\n\n用户要求：{ci}" if ci else"")
            claims={}
            try:claims=extract_json(call_llm("deepseek",ds_key,ds_model,S1,U1,max_tok=8192))
            except:
                try:claims=extract_json(call_llm("deepseek",ds_key,ds_model,S1,f"文档前段：\n\n{doc_text[:5000]}"))
                except:pass
            st.caption(f"提取:{len(claims.get('conclusions',[]))}结论,{len(claims.get('arguments',[]))}论点,{len(claims.get('evidence',[]))}论据")

            # Step 2: Search
            stat.info(f"联网检索...({gidx+1}/{total_groups})");pbar.progress(base_progress+int(group_progress*0.4))
            queries=[]
            for c in claims.get("conclusions",[])[:5]:
                t=c.get("text","")
                if len(t)>20:queries.append(t[:100])
            for e in claims.get("evidence",[])[:3]:
                if e.get("type")in("数据","引用"):
                    t=e.get("text","")
                    if len(t)>20:queries.append(t[:100])
            if ci:queries.append(ci[:100])
            queries=queries[:5]
            sr={};stxt=""
            for q in queries:
                r=search_web(q)
                if r:sr[q]=r
                for it in r:stxt+=f"\n[检索「{q}」]{it['s'][:150]}"
            sc=sum(len(v)for v in sr.values())
            st.caption(f"{len(queries)}次查询,{sc}条结果")

            # Step 3: Cross-validate
            stat.info(f"交叉验证...({gidx+1}/{total_groups})");pbar.progress(base_progress+int(group_progress*0.6))
            S2="你是学术交叉验证专家。逐条验证每个结论，给出修正后的最终结论。格式：\n##结论N\n-原文\n-验证分析\n-修正结论:[经修正后的最终表述]\n全部完成后，附录列出问题清单及严重程度。"
            U2=f"结论：{json.dumps(claims,ensure_ascii=False)[:3000]}\n检索：{stxt}\n要求：{ci if ci else'无'}\n摘要：{doc_text[:2000]}"
            group_results=[]
            for pid,pk,pm in models:
                try:
                    content=call_llm(pid,pk,pm,S2,U2)
                    group_results.append({"p":pid,"m":pm,"c":content,"e":None})
                except Exception as e:group_results.append({"p":pid,"m":pm,"c":"","e":str(e)})
            all_results.append({"doc":doc_info['fn'],"results":group_results,"search":sr})
            pbar.progress(base_progress+int(group_progress*0.8))

            # Step 4: Generate report for this doc
            mt="\n".join(f"##{r['p'].upper()}\n{r['c'][:4000]}"for r in group_results)
            S3="整合验证结果生成报告。结构：文档概况→修正后的最终结论→综合评估→附录问题清单。"
            try:report=call_llm("deepseek",ds_key,ds_model,S3,f"{mt}\n检索：{stxt}",max_tok=8192)
            except:report=f"#审校报告\n\n{mt}"
            all_reports.append({"doc":doc_info['fn'],"report":report,"results":group_results,"search":sr,"time":datetime.now().strftime("%H:%M")})
            pbar.progress(base_progress+group_progress)

        # Save to history
        for rpt in all_reports:
            st.session_state.history.append({
                "time":datetime.now().strftime("%m-%d %H:%M"),
                "title":rpt['doc'][:50],
                "report":rpt['report'],
                "results":rpt['results'],
                "search":rpt['search']
            })
        # Keep last 20
        st.session_state.history=st.session_state.history[-20:]

        st.session_state.all_results=all_results
        st.session_state.reports=all_reports
        st.session_state.done=True;pbar.progress(100);stat.success("完成");time.sleep(0.3);st.rerun()
    except Exception as e:
        import traceback;st.error(f"异常:\n{traceback.format_exc()}");st.session_state.running=False

# ── Results ──
if st.session_state.done and st.session_state.reports:
    st.markdown("---");st.header("审校结果")
    reports=st.session_state.reports
    all_results=st.session_state.all_results
    st.caption(f"共{len(reports)}篇报告")

    # If multiple reports, show selector
    if len(reports)>1:
        tabs=st.tabs([r['doc'][:30] for r in reports])
        for i,(tab,rpt) in enumerate(zip(tabs,reports)):
            with tab:
                total_search=sum(len(v)for v in rpt['search'].values())
                st.caption(f"检索:{total_search}条 | 时间:{rpt['time']}")
                st.markdown('<div class="report-box">',unsafe_allow_html=True)
                st.markdown(rpt['report'])
                st.markdown('</div>',unsafe_allow_html=True)
                ts=datetime.now().strftime("%Y%m%d_%H%M%S")
                st.download_button(f"下载{rpt['doc'][:20]}",data=rpt['report'].encode(),file_name=f"审校报告_{rpt['doc'][:20]}_{ts}.md",mime="text/markdown",key=f"dl_{i}")
                with st.expander(f"模型详情({len(rpt['results'])}个)"):
                    for rr in rpt['results']:
                        ic="OK" if not rr.get("e")else"ERR"
                        st.caption(f"{ic} {rr['p'].upper()}");st.text(rr.get("c","")[:2000])
    else:
        rpt=reports[0]
        total_search=sum(len(v)for v in rpt['search'].values())
        c1,c2,c3=st.columns(3)
        c1.metric("模型",len(rpt['results']));c2.metric("成功",sum(1 for r in rpt['results'] if not r.get("e")))
        c3.metric("检索",f"{total_search}条")
        t1,t2=st.tabs(["报告","详情"])
        with t1:
            st.markdown('<div class="report-box">',unsafe_allow_html=True)
            st.markdown(rpt['report']);st.markdown('</div>',unsafe_allow_html=True)
            ts=datetime.now().strftime("%Y%m%d_%H%M%S")
            st.download_button("下载报告",data=rpt['report'].encode(),file_name=f"审校报告_{rpt['doc'][:20]}_{ts}.md",mime="text/markdown",use_container_width=True)
        with t2:
            for rr in rpt['results']:
                with st.expander(f"{'OK' if not rr.get('e') else 'ERR'} {rr['p'].upper()}"):
                    if rr.get("e"):st.error(rr["e"])
                    else:st.text(rr["c"][:4000])
            if rpt['search']:
                with st.expander("检索详情"):
                    for q,items in rpt['search'].items():
                        st.caption(f"搜索:{q}")
                        for it in items:st.caption(it["s"][:150])
    if st.button("分析新文档",use_container_width=True):reset();st.rerun()
