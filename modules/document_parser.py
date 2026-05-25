"""
DocumentParser — parse PDF, DOCX, MD, TXT, HTML into structured text.
"""

import io
from dataclasses import dataclass
from pathlib import Path
from typing import Optional


@dataclass
class ParsedDocument:
    filename: str
    file_type: str
    text: str
    char_count: int
    page_count: int = 0
    error: Optional[str] = None


def parse_pdf(file_bytes: bytes, filename: str) -> ParsedDocument:
    """Parse PDF using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = []
        for page in doc:
            pages.append(page.get_text("text"))
        doc.close()
        text = "\n\n".join(pages)
        return ParsedDocument(
            filename=filename,
            file_type="pdf",
            text=text,
            char_count=len(text),
            page_count=len(pages),
        )
    except Exception as e:
        return ParsedDocument(
            filename=filename,
            file_type="pdf",
            text="",
            char_count=0,
            error=f"PDF解析失败: {e}",
        )


def parse_docx(file_bytes: bytes, filename: str) -> ParsedDocument:
    """Parse DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        text = "\n\n".join(paragraphs)
        return ParsedDocument(
            filename=filename,
            file_type="docx",
            text=text,
            char_count=len(text),
        )
    except Exception as e:
        return ParsedDocument(
            filename=filename,
            file_type="docx",
            text="",
            char_count=0,
            error=f"DOCX解析失败: {e}",
        )


def parse_html(file_bytes: bytes, filename: str) -> ParsedDocument:
    """Parse HTML using BeautifulSoup."""
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(file_bytes, "html.parser")
        # Remove script and style elements
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        text = soup.get_text(separator="\n", strip=True)
        return ParsedDocument(
            filename=filename,
            file_type="html",
            text=text,
            char_count=len(text),
        )
    except Exception as e:
        return ParsedDocument(
            filename=filename,
            file_type="html",
            text="",
            char_count=0,
            error=f"HTML解析失败: {e}",
        )


def parse_text(file_bytes: bytes, filename: str) -> ParsedDocument:
    """Parse plain text (MD/TXT)."""
    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        try:
            text = file_bytes.decode("gbk")
        except Exception:
            try:
                text = file_bytes.decode("latin-1")
            except Exception as e:
                return ParsedDocument(
                    filename=filename,
                    file_type="txt",
                    text="",
                    char_count=0,
                    error=f"编码识别失败: {e}",
                )

    file_type = "md" if filename.lower().endswith(".md") else "txt"
    return ParsedDocument(
        filename=filename,
        file_type=file_type,
        text=text,
        char_count=len(text),
    )


PARSERS = {
    "pdf": parse_pdf,
    "docx": parse_docx,
    "html": parse_html,
    "htm": parse_html,
    "md": parse_text,
    "txt": parse_text,
}


def parse_document(file_bytes: bytes, filename: str) -> ParsedDocument:
    """Auto-detect file type and parse."""
    ext = Path(filename).suffix.lower().lstrip(".")
    parser = PARSERS.get(ext)
    if parser is None:
        return ParsedDocument(
            filename=filename,
            file_type=ext,
            text="",
            char_count=0,
            error=f"不支持的文件格式: .{ext}（支持: PDF, DOCX, MD, TXT, HTML）",
        )
    return parser(file_bytes, filename)


def parse_multiple(uploaded_files: list) -> list[ParsedDocument]:
    """Parse multiple uploaded files. Each file is (filename, bytes)."""
    results = []
    for filename, file_bytes in uploaded_files:
        results.append(parse_document(file_bytes, filename))
    return results
